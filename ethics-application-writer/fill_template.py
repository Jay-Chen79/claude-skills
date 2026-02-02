#!/usr/bin/env python3
"""
伦理申请材料模板填充工具
保留原Word模板格式，仅替换占位符内容
"""

import sys
import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("需要安装 python-docx 库:")
    print("  pip install python-docx")
    sys.exit(1)


def replace_in_paragraph(paragraph, replacements: dict):
    """
    在段落中替换文本，保留原有格式
    """
    for key, value in replacements.items():
        if key in paragraph.text:
            # 遍历所有run，保留格式替换
            for run in paragraph.runs:
                if key in run.text:
                    run.text = run.text.replace(key, str(value))


def replace_in_table(table, replacements: dict):
    """
    在表格中替换文本，保留原有格式
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, replacements)


def fill_template(template_path: str, output_path: str, replacements: dict):
    """
    填充Word模板

    Args:
        template_path: 模板文件路径
        output_path: 输出文件路径
        replacements: 替换字典 {"【占位符】": "实际内容"}
    """
    # 读取模板
    doc = Document(template_path)

    # 替换段落中的内容
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)

    # 替换表格中的内容
    for table in doc.tables:
        replace_in_table(table, replacements)

    # 替换页眉中的内容
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            replace_in_paragraph(paragraph, replacements)

        footer = section.footer
        for paragraph in footer.paragraphs:
            replace_in_paragraph(paragraph, replacements)

    # 保存文件
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)

    doc.save(output_path)
    print(f"文件已保存: {output_path}")


# 模板目录
TEMPLATE_DIR = "/Users/chenzheyi/Desktop/Document_M1/涉及人的生物医学伦理审查申请"

# 模板文件映射
TEMPLATES = {
    "初始审查申请表": "初始审查申请表.docx",
    "研究方案": "研究方案.doc",
    "知情同意书_干预": "知情同意 干预.docx",
    "知情同意书_非干预": "知情同意 非干预.docx",
    "知情同意书_监护人": "知情同意 监护人.docx",
    "知情同意书_未成年人": "知情同意 未成年人.docx",
    "研究者简历": "研究者简历.docx",
    "利益冲突声明": "研究者利益冲突声明.docx",
    "PI责任声明": "主要研究者责任声明.docx",
    "自查清单": "自查清单.docx",
    "递交信": "伦理审查文件递交信.docx",
    "备案表": "备案材料/2.临床研究项目备案表 5.2.docx",
    "团队名单": "备案材料/研究团队名单.docx",
    "报送资料列表": "备案材料/临床研究项目报送资料列表3.0.docx",
}


def get_template_path(template_name: str) -> str:
    """获取模板完整路径"""
    if template_name in TEMPLATES:
        return os.path.join(TEMPLATE_DIR, TEMPLATES[template_name])
    raise ValueError(f"未知模板: {template_name}")


def generate_output_path(project_name: str, template_name: str, version: str = "1.0") -> str:
    """生成输出文件路径"""
    from datetime import datetime
    date_str = datetime.now().strftime("%Y%m%d")
    filename = f"{project_name}_{template_name}_V{version}_{date_str}.docx"
    return os.path.join(TEMPLATE_DIR, "备案材料", filename)


# 常用占位符示例
COMMON_PLACEHOLDERS = {
    # 项目信息
    "【项目名称】": "",
    "【项目编号】": "",
    "【研究类型】": "",
    "【资金来源】": "",

    # 研究者信息
    "【主要研究者】": "",
    "【职称】": "",
    "【科室】": "",
    "【联系电话】": "",
    "【电子邮箱】": "",

    # 研究设计
    "【研究目的】": "",
    "【样本量】": "",
    "【研究周期】": "",
    "【开始日期】": "",
    "【结束日期】": "",

    # 受试者信息
    "【纳入标准】": "",
    "【排除标准】": "",
    "【风险描述】": "",
    "【获益描述】": "",
}


if __name__ == "__main__":
    # 使用示例
    print("伦理申请材料模板填充工具")
    print("=" * 50)
    print("\n可用模板:")
    for name, path in TEMPLATES.items():
        print(f"  - {name}: {path}")

    print("\n使用方法:")
    print("""
    from fill_template import fill_template, get_template_path, generate_output_path

    # 定义替换内容
    replacements = {
        "【项目名称】": "成人框架眼镜试戴体验评估量表的研制与验证",
        "【主要研究者】": "陈浙一",
        "【样本量】": "500",
        # ... 其他字段
    }

    # 填充模板
    template = get_template_path("初始审查申请表")
    output = generate_output_path("眼镜试戴量表", "初始审查申请表")
    fill_template(template, output, replacements)
    """)
