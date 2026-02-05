#!/usr/bin/env python3
"""
参考文献解析模块：提取并分析docx文件中的引用和参考文献
增强版：提取引用上下文用于相关性分析
"""

import re
import json
from dataclasses import dataclass, asdict, field
from typing import List, Dict, Tuple, Optional, Set
from docx import Document


@dataclass
class Citation:
    """引用标记"""
    number: int
    text: str  # 引用标记文本，如 [1]
    position: int  # 在段落中的位置
    para_index: int  # 段落索引
    context: str = ""  # 引用上下文（前后句子）


@dataclass
class Reference:
    """参考文献条目"""
    number: int
    raw_text: str  # 原始文本
    ref_type: str  # journal/book/conference/web/thesis/other
    title: str = ""
    authors: List[str] = field(default_factory=list)
    year: str = ""
    journal: str = ""
    doi: str = ""
    pmid: str = ""
    url: str = ""
    isbn: str = ""
    is_valid: bool = True
    issues: List[str] = field(default_factory=list)


@dataclass
class ReferenceCheckResult:
    """解析结果"""
    citations: List[Citation]
    references: List[Reference]
    missing_refs: List[int]  # 文中引用但列表无
    unused_refs: List[int]  # 列表有但文中无
    out_of_order: List[Tuple[int, int]]  # (期望, 实际)
    duplicates: List[int]  # 重复引用


class ReferenceParser:
    """参考文献解析器"""

    # 引用标记模式
    CITATION_PATTERNS = [
        r"\[(\d+)\]",  # [1]
        r"\[(\d+)-(\d+)\]",  # [1-3]
        r"\[(\d+)(?:,\s*\d+)+\]",  # [1,2,3]
    ]

    # 参考文献章节标记
    REF_SECTION_MARKERS = [
        r"^参考文献\s*$",
        r"^References\s*$",
        r"^REFERENCES\s*$",
        r"^Bibliography\s*$",
        r"^文\s*献\s*$",
        r"^引用文献\s*$",
    ]

    def __init__(self, docx_path: str):
        self.doc = Document(docx_path)
        self.paragraphs_text: List[str] = []
        self.body_paragraphs: List[str] = []
        self.ref_section_text = ""
        self.ref_section_start_idx = -1
        self._extract_text()

    def _extract_text(self):
        """提取文档文本，分离正文和参考文献部分"""
        ref_section_started = False

        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            self.paragraphs_text.append(text)

            if not text:
                continue

            # 检查是否进入参考文献章节
            if not ref_section_started:
                for marker in self.REF_SECTION_MARKERS:
                    if re.match(marker, text, re.IGNORECASE):
                        ref_section_started = True
                        self.ref_section_start_idx = i
                        break

            if ref_section_started:
                self.ref_section_text += text + "\n"
            else:
                self.body_paragraphs.append(text)

    def _get_context(self, para_idx: int, position: int, window: int = 100) -> str:
        """获取引用上下文"""
        if para_idx >= len(self.paragraphs_text):
            return ""

        text = self.paragraphs_text[para_idx]
        start = max(0, position - window)
        end = min(len(text), position + window)

        context = text[start:end]

        # 添加前后段落的部分内容
        if para_idx > 0 and self.paragraphs_text[para_idx - 1]:
            prev = self.paragraphs_text[para_idx - 1]
            context = prev[-50:] + " ... " + context

        if para_idx < len(self.paragraphs_text) - 1 and self.paragraphs_text[para_idx + 1]:
            next_para = self.paragraphs_text[para_idx + 1]
            context = context + " ... " + next_para[:50]

        return context.strip()

    def extract_citations(self) -> List[Citation]:
        """提取正文中的所有引用"""
        citations = []

        for para_idx, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # 跳过参考文献章节
            if self.ref_section_start_idx >= 0 and para_idx >= self.ref_section_start_idx:
                continue

            # 查找单个引用 [n]
            for match in re.finditer(r"\[(\d+)\]", text):
                try:
                    num = int(match.group(1))
                    context = self._get_context(para_idx, match.start())
                    citations.append(Citation(
                        number=num,
                        text=match.group(0),
                        position=match.start(),
                        para_index=para_idx,
                        context=context
                    ))
                except ValueError:
                    continue

            # 查找范围引用 [n-m]
            for match in re.finditer(r"\[(\d+)-(\d+)\]", text):
                try:
                    start_num = int(match.group(1))
                    end_num = int(match.group(2))
                    context = self._get_context(para_idx, match.start())
                    for num in range(start_num, end_num + 1):
                        citations.append(Citation(
                            number=num,
                            text=match.group(0),
                            position=match.start(),
                            para_index=para_idx,
                            context=context
                        ))
                except ValueError:
                    continue

            # 查找列表引用 [n,m,...]
            for match in re.finditer(r"\[((\d+)(?:,\s*\d+)+)\]", text):
                try:
                    nums_str = match.group(1)
                    nums = [int(n.strip()) for n in nums_str.split(",")]
                    context = self._get_context(para_idx, match.start())
                    for num in nums:
                        citations.append(Citation(
                            number=num,
                            text=match.group(0),
                            position=match.start(),
                            para_index=para_idx,
                            context=context
                        ))
                except ValueError:
                    continue

        # 按出现顺序排序
        citations.sort(key=lambda x: (x.para_index, x.position))
        return citations

    def extract_references(self) -> List[Reference]:
        """提取参考文献列表"""
        references = []

        if not self.ref_section_text:
            return references

        lines = self.ref_section_text.split("\n")
        current_ref: Optional[Reference] = None
        current_text = ""

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 检查是否是新的参考文献条目
            match = re.match(r"^\[(\d+)\]\s*(.+)", line)
            if match:
                # 保存上一条
                if current_ref is not None:
                    current_ref.raw_text = current_text.strip()
                    self._parse_reference_details(current_ref)
                    references.append(current_ref)

                try:
                    num = int(match.group(1))
                    current_ref = Reference(number=num, raw_text="", ref_type="other")
                    current_text = match.group(2)
                except ValueError:
                    current_ref = None
            else:
                # 续行
                if current_ref is not None:
                    current_text += " " + line

        # 保存最后一条
        if current_ref is not None:
            current_ref.raw_text = current_text.strip()
            self._parse_reference_details(current_ref)
            references.append(current_ref)

        return references

    def _parse_reference_details(self, ref: Reference):
        """解析参考文献详细信息"""
        text = ref.raw_text

        # 判断类型
        if re.search(r"\[J\]", text) or re.search(r"\d{4}[a-z]?[;,]\s*\d+\s*\(\d+\)", text):
            ref.ref_type = "journal"
        elif re.search(r"\[M\]", text) or re.search(r"(?:出版社|Press|Publishing|Publishers)", text, re.IGNORECASE):
            ref.ref_type = "book"
        elif re.search(r"\[C\]", text) or re.search(r"(?:会议|Conference|Proceedings|Symposium)", text, re.IGNORECASE):
            ref.ref_type = "conference"
        elif re.search(r"\[D\]", text) or re.search(r"(?:博士|硕士|学位|Thesis|Dissertation)", text, re.IGNORECASE):
            ref.ref_type = "thesis"
        elif re.search(r"\[EB/OL\]", text) or re.search(r"https?://", text):
            ref.ref_type = "web"
        else:
            ref.ref_type = "other"

        # 提取 DOI
        doi_match = re.search(r"(?:doi[:\s]*)?10\.\d{4,}/[^\s\]]+", text, re.IGNORECASE)
        if doi_match:
            doi = doi_match.group(0)
            # 清理 DOI
            if doi.lower().startswith("doi"):
                doi = re.sub(r"^doi[:\s]*", "", doi, flags=re.IGNORECASE)
            ref.doi = doi.rstrip(".,;")

        # 提取 PMID
        pmid_match = re.search(r"(?:PMID|PubMed\s*(?:ID)?)[:\s]*(\d{1,8})", text, re.IGNORECASE)
        if pmid_match:
            ref.pmid = pmid_match.group(1)

        # 提取 URL
        url_match = re.search(r"(https?://[^\s\]]+)", text)
        if url_match:
            ref.url = url_match.group(1).rstrip(".,;")

        # 提取 ISBN
        isbn_match = re.search(r"ISBN[:\s]*([\d\-X]+)", text, re.IGNORECASE)
        if isbn_match:
            ref.isbn = isbn_match.group(1)

        # 提取年份
        year_match = re.search(r"(19|20)\d{2}", text)
        if year_match:
            ref.year = year_match.group(0)

        # 提取标题（尝试多种模式）
        # GB/T 7714 格式：作者. 标题[J/M/C/D]. ...
        title_match = re.search(r"\.\s*([^.\[\]]{10,}?)\s*\[[JMCDEB/OL]+\]", text)
        if title_match:
            ref.title = title_match.group(1).strip()
        else:
            # APA/其他格式：尝试提取引号内的标题或第一个句号后的内容
            title_match = re.search(r'"([^"]+)"', text)
            if title_match:
                ref.title = title_match.group(1).strip()
            else:
                # 提取第一个句号后到下一个句号的内容
                parts = text.split(".")
                if len(parts) > 1:
                    ref.title = parts[1].strip()[:100]

    def check_citation_order(self, citations: List[Citation], references: List[Reference]) -> ReferenceCheckResult:
        """检查引用顺序和完整性"""
        result = ReferenceCheckResult(
            citations=citations,
            references=references,
            missing_refs=[],
            unused_refs=[],
            out_of_order=[],
            duplicates=[],
        )

        # 收集引用的序号
        cited_numbers: Set[int] = set()
        for c in citations:
            cited_numbers.add(c.number)

        # 收集参考文献的序号
        ref_numbers = {r.number for r in references}

        # 检查缺失和未使用
        result.missing_refs = sorted(list(cited_numbers - ref_numbers))
        result.unused_refs = sorted(list(ref_numbers - cited_numbers))

        # 检查重复（同一序号多次出现是正常的，但需要记录）
        seen_first: Dict[int, int] = {}  # 序号 -> 首次出现的索引
        for i, c in enumerate(citations):
            if c.number not in seen_first:
                seen_first[c.number] = i

        # 检查顺序：首次出现应该按递增顺序
        first_occurrences = [(num, idx) for num, idx in seen_first.items()]
        first_occurrences.sort(key=lambda x: x[1])

        expected_num = 1
        for num, _ in first_occurrences:
            if num != expected_num:
                if num > expected_num:
                    result.out_of_order.append((expected_num, num))
                expected_num = num + 1
            else:
                expected_num += 1

        return result

    def parse(self) -> ReferenceCheckResult:
        """执行完整解析"""
        citations = self.extract_citations()
        references = self.extract_references()
        result = self.check_citation_order(citations, references)
        return result

    def export_to_json(self, result: ReferenceCheckResult, output_path: str):
        """导出结果到 JSON"""
        data = {
            "citations": [asdict(c) for c in result.citations],
            "references": [asdict(r) for r in result.references],
            "missing_refs": result.missing_refs,
            "unused_refs": result.unused_refs,
            "out_of_order": result.out_of_order,
            "duplicates": result.duplicates,
            "summary": {
                "total_citations": len(result.citations),
                "total_references": len(result.references),
                "unique_cited": len(set(c.number for c in result.citations)),
                "missing_count": len(result.missing_refs),
                "unused_count": len(result.unused_refs),
                "order_issues": len(result.out_of_order),
            }
        }

        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)


def main():
    import sys

    if len(sys.argv) < 2:
        print("Usage: python parse_references.py <docx_file> [output_json]")
        print("\nExample:")
        print("  python parse_references.py paper.docx parsed.json")
        sys.exit(1)

    docx_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else "references_parsed.json"

    try:
        parser = ReferenceParser(docx_path)
        result = parser.parse()
        parser.export_to_json(result, output_path)

        print(f"解析完成！")
        print(f"  - 文中引用数: {len(result.citations)}")
        print(f"  - 唯一引用数: {len(set(c.number for c in result.citations))}")
        print(f"  - 参考文献数: {len(result.references)}")
        print(f"  - 缺失的参考文献: {result.missing_refs or '无'}")
        print(f"  - 未被引用的参考文献: {result.unused_refs or '无'}")
        print(f"  - 顺序错误: {result.out_of_order or '无'}")
        print(f"\n结果已保存到: {output_path}")

    except Exception as e:
        print(f"错误: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
