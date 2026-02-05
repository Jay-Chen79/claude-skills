#!/usr/bin/env python3
"""
生成检查报告和修改后的docx文件
增强版：符合详细报告格式要求
"""

import json
import re
from datetime import datetime
from typing import Dict, List, Any, Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT


class ReportGenerator:
    """报告生成器"""

    def __init__(self, original_docx: str, parsed_json: str, verified_json: str):
        self.original_docx = original_docx
        self.parsed_json = parsed_json
        self.verified_json = verified_json

        with open(parsed_json, "r", encoding="utf-8") as f:
            self.parsed_data = json.load(f)

        with open(verified_json, "r", encoding="utf-8") as f:
            self.verified_data = json.load(f)

    def _get_severity(self) -> str:
        """评估紧急程度"""
        invalid_count = self.verified_data.get("summary", {}).get("invalid", 0)
        suspicious_count = self.verified_data.get("summary", {}).get("suspicious", 0)
        missing_refs = len(self.parsed_data.get("missing_refs", []))
        order_issues = len(self.parsed_data.get("out_of_order", []))

        if invalid_count > 0 or missing_refs > 3:
            return "高"
        elif suspicious_count > 2 or missing_refs > 0 or order_issues > 3:
            return "中"
        else:
            return "低"

    def _get_overall_assessment(self) -> str:
        """总体评估"""
        validity_rate = self.verified_data.get("summary", {}).get("validity_rate", "0%")
        rate = float(validity_rate.replace("%", ""))

        missing_refs = len(self.parsed_data.get("missing_refs", []))
        order_issues = len(self.parsed_data.get("out_of_order", []))

        if rate >= 95 and missing_refs == 0 and order_issues == 0:
            return "优秀"
        elif rate >= 80 and missing_refs <= 1:
            return "良好"
        elif rate >= 60:
            return "需修改"
        else:
            return "存在严重问题"

    def generate_report_docx(self, output_path: str):
        """生成详细检查报告"""
        doc = Document()

        # 设置页边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)

        # 标题
        title = doc.add_heading("参考文献检查报告", 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 生成时间
        doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()

        # ========== A. 执行摘要 ==========
        doc.add_heading("A. 执行摘要", 1)

        summary = self.verified_data.get("summary", {})
        total = summary.get("total", 0)
        valid = summary.get("valid", 0)
        invalid = summary.get("invalid", 0)
        suspicious = summary.get("suspicious", 0)
        validity_rate = summary.get("validity_rate", "0%")

        parsed_summary = self.parsed_data.get("summary", {})
        missing_count = len(self.parsed_data.get("missing_refs", []))
        unused_count = len(self.parsed_data.get("unused_refs", []))
        order_issues = len(self.parsed_data.get("out_of_order", []))

        # 总体评估
        p = doc.add_paragraph()
        p.add_run("总体评估：").bold = True
        p.add_run(self._get_overall_assessment())

        # 统计表
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("统计数据：").bold = True

        table = doc.add_table(rows=7, cols=2)
        table.style = "Light Grid Accent 1"

        stats = [
            ("参考文献总数", str(total)),
            ("验证有效", f"{valid} ({validity_rate})"),
            ("存疑文献", str(suspicious)),
            ("无法验证", str(invalid)),
            ("格式错误", str(missing_count + unused_count)),
            ("序号问题", str(order_issues)),
            ("紧急程度", self._get_severity()),
        ]

        for i, (key, value) in enumerate(stats):
            table.rows[i].cells[0].text = key
            table.rows[i].cells[1].text = value

        # 紧急程度说明
        doc.add_paragraph()
        severity = self._get_severity()
        if severity == "高":
            p = doc.add_paragraph()
            run = p.add_run("⚠️ 紧急程度：高 - 存在虚假文献或严重格式错误，需立即处理")
            run.font.color.rgb = RGBColor(255, 0, 0)
        elif severity == "中":
            p = doc.add_paragraph()
            run = p.add_run("⚠️ 紧急程度：中 - 部分文献无法验证或有格式问题，建议修改")
            run.font.color.rgb = RGBColor(255, 165, 0)
        else:
            p = doc.add_paragraph()
            run = p.add_run("✅ 紧急程度：低 - 仅有轻微问题或无问题")
            run.font.color.rgb = RGBColor(0, 128, 0)

        doc.add_paragraph()

        # ========== B. 详细检查结果 ==========
        doc.add_heading("B. 详细检查结果", 1)

        results = self.verified_data.get("results", [])

        if results:
            # 创建表格
            table = doc.add_table(rows=1, cols=6)
            table.style = "Light Grid Accent 1"

            # 表头
            headers = ["序号", "原文献信息", "检查状态", "DOI/PubMed链接", "相关性评分", "问题说明"]
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = header
                table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

            # 数据行
            for result in results:
                row = table.add_row()
                row.cells[0].text = f"[{result.get('reference_number', '?')}]"

                # 文献信息摘要
                raw_text = result.get("raw_text", "")
                row.cells[1].text = raw_text[:80] + "..." if len(raw_text) > 80 else raw_text

                # 检查状态
                status = result.get("status", "unknown")
                status_text = {
                    "verified": "✅ 已验证",
                    "unverified": "❓ 无法验证",
                    "suspicious": "⚠️ 存疑",
                    "invalid": "❌ 无效",
                }.get(status, status)
                row.cells[2].text = status_text

                # 链接
                verified_url = result.get("verified_url", "")
                row.cells[3].text = verified_url[:50] + "..." if len(verified_url) > 50 else verified_url

                # 相关性评分
                relevance = result.get("relevance", "待审核")
                row.cells[4].text = relevance

                # 问题说明
                issues = result.get("issues", [])
                row.cells[5].text = "; ".join(issues) if issues else "无"

        else:
            doc.add_paragraph("暂无验证结果数据。")

        doc.add_paragraph()

        # ========== C. 问题分类清单 ==========
        doc.add_heading("C. 问题分类清单", 1)

        # 虚假/不存在的文献
        invalid_refs = [r for r in results if r.get("status") == "invalid"]
        if invalid_refs:
            doc.add_heading("虚假/不存在的文献", 2)
            for ref in invalid_refs:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"[{ref.get('reference_number', '?')}] ").bold = True
                p.add_run(f"{ref.get('raw_text', '')[:100]}...")
                if ref.get("issues"):
                    p.add_run(f"\n  问题：{'; '.join(ref['issues'])}")

        # 无法验证的文献
        unverified_refs = [r for r in results if r.get("status") == "unverified"]
        if unverified_refs:
            doc.add_heading("无法验证的文献", 2)
            for ref in unverified_refs:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"[{ref.get('reference_number', '?')}] ").bold = True
                p.add_run(f"{ref.get('raw_text', '')[:100]}...")

        # 序号顺序问题
        out_of_order = self.parsed_data.get("out_of_order", [])
        if out_of_order:
            doc.add_heading("序号顺序问题", 2)
            for expected, actual in out_of_order:
                doc.add_paragraph(f"期望 [{expected}]，实际 [{actual}]", style="List Bullet")

        # 格式错误
        missing_refs = self.parsed_data.get("missing_refs", [])
        unused_refs = self.parsed_data.get("unused_refs", [])
        if missing_refs or unused_refs:
            doc.add_heading("格式错误", 2)
            if missing_refs:
                doc.add_paragraph(f"缺失的参考文献（文中引用但列表无）：{missing_refs}", style="List Bullet")
            if unused_refs:
                doc.add_paragraph(f"未被引用的参考文献（列表有但文中无）：{unused_refs}", style="List Bullet")

        # 相关性存疑
        suspicious_relevance = [r for r in results if r.get("relevance") in ["相关性存疑", "明显不相关"]]
        if suspicious_relevance:
            doc.add_heading("相关性存疑", 2)
            for ref in suspicious_relevance:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"[{ref.get('reference_number', '?')}] ").bold = True
                p.add_run(f"相关性：{ref.get('relevance', '?')}")

        if not any([invalid_refs, unverified_refs, out_of_order, missing_refs, unused_refs, suspicious_relevance]):
            doc.add_paragraph("✅ 未发现明显问题。")

        doc.add_paragraph()

        # ========== D. 修改建议 ==========
        doc.add_heading("D. 修改建议", 1)

        suggestions = []

        if missing_refs:
            suggestions.append(f"补充参考文献列表中缺失的条目：{missing_refs}")

        if unused_refs:
            suggestions.append(f"删除未被引用的参考文献或在正文中添加引用：{unused_refs}")

        if out_of_order:
            suggestions.append("调整参考文献编号，使其按文中首次出现顺序排列")

        if invalid_refs:
            invalid_nums = [r.get("reference_number") for r in invalid_refs]
            suggestions.append(f"核实并替换以下无效/虚假文献：{invalid_nums}")

        if unverified_refs:
            unverified_nums = [r.get("reference_number") for r in unverified_refs]
            suggestions.append(f"补充以下文献的 DOI 或 PubMed ID 以便验证：{unverified_nums}")

        if suspicious_relevance:
            susp_nums = [r.get("reference_number") for r in suspicious_relevance]
            suggestions.append(f"人工审核以下文献与引用上下文的相关性：{susp_nums}")

        if suggestions:
            for i, suggestion in enumerate(suggestions, 1):
                doc.add_paragraph(f"{i}. {suggestion}")
        else:
            doc.add_paragraph("✅ 未发现需要修改的问题。")

        doc.add_paragraph()

        # ========== E. 验证链接汇总 ==========
        doc.add_heading("E. 验证链接汇总", 1)

        verified_with_links = [r for r in results if r.get("verified_url")]
        if verified_with_links:
            # DOI 链接
            doi_links = [r for r in verified_with_links if "doi.org" in r.get("verified_url", "")]
            if doi_links:
                doc.add_heading("DOI 链接", 2)
                for ref in doi_links:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(f"[{ref.get('reference_number', '?')}] ")
                    p.add_run(ref.get("verified_url", ""))

            # PubMed 链接
            pubmed_links = [r for r in verified_with_links if "pubmed" in r.get("verified_url", "").lower()]
            if pubmed_links:
                doc.add_heading("PubMed 链接", 2)
                for ref in pubmed_links:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(f"[{ref.get('reference_number', '?')}] ")
                    p.add_run(ref.get("verified_url", ""))

            # 其他链接
            other_links = [r for r in verified_with_links if
                           "doi.org" not in r.get("verified_url", "") and
                           "pubmed" not in r.get("verified_url", "").lower()]
            if other_links:
                doc.add_heading("其他验证链接", 2)
                for ref in other_links:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(f"[{ref.get('reference_number', '?')}] ")
                    p.add_run(ref.get("verified_url", ""))
        else:
            doc.add_paragraph("暂无验证链接。")

        doc.add_paragraph()

        # 页脚
        footer = doc.add_paragraph()
        footer.add_run("— 报告结束 —").italic = True
        footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.save(output_path)
        print(f"检查报告已生成：{output_path}")

    def generate_corrected_docx(self, output_path: str):
        """生成标注问题的文档副本"""
        doc = Document(self.original_docx)

        # 构建验证结果映射
        verified_results: Dict[int, Dict] = {}
        for r in self.verified_data.get("results", []):
            verified_results[r.get("reference_number", 0)] = r

        # 标记问题序号集合
        problem_refs = set()
        for r in self.verified_data.get("results", []):
            if r.get("status") in ["invalid", "suspicious", "unverified"]:
                problem_refs.add(r.get("reference_number", 0))

        # 添加缺失和未使用的也标记
        problem_refs.update(self.parsed_data.get("missing_refs", []))

        # 遍历正文，标记问题引用
        ref_section_started = False
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 检查是否进入参考文献章节
            if re.match(r"^(参考文献|References|REFERENCES|Bibliography)", text, re.IGNORECASE):
                ref_section_started = True
                continue

            if not ref_section_started:
                # 正文部分：标记问题引用为红色
                for run in para.runs:
                    run_text = run.text
                    # 查找引用标记
                    citations = re.findall(r"\[(\d+)\]", run_text)
                    for cite_num in citations:
                        try:
                            num = int(cite_num)
                            if num in problem_refs:
                                run.font.color.rgb = RGBColor(255, 0, 0)
                                break
                        except ValueError:
                            continue
            else:
                # 参考文献部分：标记问题条目
                match = re.match(r"^\[(\d+)\]", text)
                if match:
                    try:
                        num = int(match.group(1))
                        if num in problem_refs:
                            for run in para.runs:
                                run.font.color.rgb = RGBColor(255, 0, 0)
                                run.font.bold = True
                    except ValueError:
                        pass

        doc.save(output_path)
        print(f"标注后的文档已生成：{output_path}")

    def generate_correction_list(self, output_path: str):
        """生成机器可读的修改清单"""
        corrections = {
            "generated_at": datetime.now().isoformat(),
            "issues": []
        }

        # 缺失的参考文献
        missing_refs = self.parsed_data.get("missing_refs", [])
        if missing_refs:
            corrections["issues"].append({
                "type": "missing_references",
                "description": "文中引用但参考文献列表中缺失",
                "severity": "high",
                "items": missing_refs,
                "action": "补充参考文献或删除对应引用"
            })

        # 未被引用的参考文献
        unused_refs = self.parsed_data.get("unused_refs", [])
        if unused_refs:
            corrections["issues"].append({
                "type": "unused_references",
                "description": "参考文献列表中存在但文中未引用",
                "severity": "medium",
                "items": unused_refs,
                "action": "删除未使用的参考文献或在正文中添加引用"
            })

        # 顺序错误
        out_of_order = self.parsed_data.get("out_of_order", [])
        if out_of_order:
            corrections["issues"].append({
                "type": "order_issues",
                "description": "引用顺序错误",
                "severity": "medium",
                "items": [{"expected": e, "actual": a} for e, a in out_of_order],
                "action": "按文中首次出现顺序重新编号"
            })

        # 无效/存疑的参考文献
        results = self.verified_data.get("results", [])
        invalid_refs = [r for r in results if r.get("status") == "invalid"]
        if invalid_refs:
            corrections["issues"].append({
                "type": "invalid_references",
                "description": "无法验证真实性的参考文献",
                "severity": "high",
                "items": [{
                    "number": r.get("reference_number"),
                    "text": r.get("raw_text", "")[:150],
                    "issues": r.get("issues", [])
                } for r in invalid_refs],
                "action": "核实并替换虚假文献"
            })

        suspicious_refs = [r for r in results if r.get("status") == "suspicious"]
        if suspicious_refs:
            corrections["issues"].append({
                "type": "suspicious_references",
                "description": "存疑的参考文献（需人工确认）",
                "severity": "medium",
                "items": [{
                    "number": r.get("reference_number"),
                    "text": r.get("raw_text", "")[:150],
                    "issues": r.get("issues", [])
                } for r in suspicious_refs],
                "action": "人工核实文献真实性"
            })

        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(corrections, f, ensure_ascii=False, indent=2)

        print(f"修改清单已生成：{output_path}")


def main():
    import sys

    if len(sys.argv) < 4:
        print("Usage: python generate_report.py <original.docx> <parsed.json> <verified.json> [output_prefix]")
        print("\nExample:")
        print("  python generate_report.py paper.docx parsed.json verified.json reference_check")
        print("\nOutput files:")
        print("  reference_check_report.docx    - 详细检查报告")
        print("  reference_check_corrected.docx - 标注问题的文档副本")
        print("  reference_check_corrections.json - 机器可读的修改清单")
        sys.exit(1)

    original_docx = sys.argv[1]
    parsed_json = sys.argv[2]
    verified_json = sys.argv[3]
    output_prefix = sys.argv[4] if len(sys.argv) > 4 else "reference_check"

    try:
        generator = ReportGenerator(original_docx, parsed_json, verified_json)

        report_path = f"{output_prefix}_report.docx"
        corrected_path = f"{output_prefix}_corrected.docx"
        corrections_path = f"{output_prefix}_corrections.json"

        generator.generate_report_docx(report_path)
        generator.generate_corrected_docx(corrected_path)
        generator.generate_correction_list(corrections_path)

        print("\n所有文件已生成完成！")

    except Exception as e:
        print(f"错误: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
